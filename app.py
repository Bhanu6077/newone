from flask import Flask, render_template, request, send_file, after_this_request
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from copy import deepcopy
import os
import re
from lxml import etree
import uuid
import logging


# ==========================================================
# APP SETUP
# ==========================================================
app = Flask(__name__)
app.config.update(
    UPLOAD_FOLDER="uploads",
    OUTPUT_FOLDER="output",
    TEMPLATE_PATH="template.docx",
    MAX_CONTENT_LENGTH=50 * 1024 * 1024,  # 50MB
    MAX_FILENAME_LENGTH=80,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
log = logging.getLogger(__name__)

ZONE_COLORS = {
    "Zone A": RGBColor(255, 0, 0),
    "Zone B": RGBColor(0, 0, 255),
    "Zone C": RGBColor(0, 128, 0),
    "Zone D": RGBColor(255, 255, 0),
    "Zone E": RGBColor(128, 0, 128),
}


# ==========================================================
# FILE HELPERS
# ==========================================================
def ensure_folders():
    for key in ("UPLOAD_FOLDER", "OUTPUT_FOLDER"):
        path = app.config[key]
        os.makedirs(path, exist_ok=True)


def sanitize_filename(name: str) -> str:
    """Remove illegal chars and trim to safe length."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    name = name.strip(". ")
    max_len = app.config["MAX_FILENAME_LENGTH"]
    return name[:max_len] if len(name) > max_len else name or "Report"


def save_upload(file, prefix: str) -> str:
    """Save an uploaded file to UPLOAD_FOLDER with a unique name. Returns path."""
    folder = app.config["UPLOAD_FOLDER"]
    ext = os.path.splitext(file.filename)[1] or ".docx"
    dest = os.path.join(folder, f"{prefix}_{uuid.uuid4().hex}{ext}")
    file.save(dest)
    return dest


def cleanup(*paths):
    """Silently delete a list of file paths."""
    for p in paths:
        try:
            if p and os.path.exists(p):
                os.remove(p)
        except Exception as e:
            log.warning("Cleanup failed for %s: %s", p, e)


# ==========================================================
# XML ITERATION HELPERS
# ==========================================================
def _block_text(block) -> str:
    return "".join(n.text for n in block.iter() if n.text).strip()


def _iter_body(source_doc):
    """Yield (tag_suffix, element) for each direct child of body."""
    for block in source_doc.element.body:
        tag = block.tag.split("}")[-1]   # 'p', 'tbl', 'sectPr', etc.
        yield tag, block


def _find_marker(target_doc, marker_text):
    """
    Find the marker paragraph in target_doc.
    Returns (parent, index, paragraph_element) or (None, None, None).
    """
    for para in target_doc.paragraphs:
        if marker_text in para.text:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            return parent, idx, para._element
    log.warning("Marker not found: %s", marker_text)
    return None, None, None


# ==========================================================
# ROAD TABLE
# ==========================================================
def insert_road_table(doc, marker_text: str, road_data: dict):
    """
    Insert a road inventory table at marker_text position.
    road_data = {"Service Roads": {"LHS": 2, "RHS": 3}, ...}
    """
    parent, index, marker_el = _find_marker(doc, marker_text)
    if parent is None:
        return

    parent.remove(marker_el)

    max_rows = max((v["LHS"] + v["RHS"] for v in road_data.values()), default=0)
    if max_rows == 0:
        return

    num_types = len(road_data)
    table = doc.add_table(rows=max_rows + 1, cols=num_types)
    table.style = "Table Grid"

    for col, road_type in enumerate(road_data.keys()):
        cell = table.cell(0, col)
        cell.text = road_type
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for col, (road_type, sides) in enumerate(road_data.items()):
        abbr = "".join(w[0] for w in road_type.split()).upper()
        display = road_type.rstrip("s")
        row_index = 1
        for i in range(1, sides["LHS"] + 1):
            table.cell(row_index, col).text = f"{abbr}L {i} – {display} LHS {i}"
            row_index += 1
        for i in range(1, sides["RHS"] + 1):
            table.cell(row_index, col).text = f"{abbr}R {i} – {display} RHS {i}"
            row_index += 1

    parent.insert(index, table._element)
    log.info("Road table inserted (%d types, %d rows)", num_types, max_rows)


# ==========================================================
# EXECUTIVE SUMMARY TABLE COPY
# ==========================================================
def copy_executive_summary_table(source_doc, target_doc, marker_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    parent.remove(marker_el)
    found_ref = False

    for tag, block in _iter_body(source_doc):
        if tag == "p":
            if "summary of Gap study report" in _block_text(block).lower():
                found_ref = True
                continue
        if found_ref and tag == "tbl":
            parent.insert(index, deepcopy(block))
            log.info("Executive summary table copied.")
            return

    log.warning("Executive summary table not found in source document.")


# ==========================================================
# TABLE FORMATTING UTILITIES
# ==========================================================
def fix_table_width(table, width_cm: float = 16.0):
    """Set table to fixed width."""
    table.autofit = False
    num_cols = len(table.columns)
    if num_cols == 0:
        return
    col_width = Cm(width_cm) / num_cols
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tblPr.append(tblW)


def fix_table_layout(table):
    """Prevent rows from splitting across pages."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        cant.set(qn("w:val"), "1")
        trPr.append(cant)


def reduce_table_font(table, size_pt: int = 9):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)


def format_table(table):
    """Apply all standard formatting to a table."""
    fix_table_width(table)
    fix_table_layout(table)
    reduce_table_font(table)


# ==========================================================
# GENERIC SECTION EXTRACTION  (paragraphs + tables + images)
# ==========================================================
def insert_section(target_doc, source_doc, marker_text: str,
                   start_heading: str, stop_heading: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    parent.remove(marker_el)
    capture = False

    for tag, block in _iter_body(source_doc):
        if tag == "p":
            text = _block_text(block)
            if start_heading in text:
                capture = True
                continue
            if capture and stop_heading and stop_heading in text:
                break

        if not capture:
            continue

        if tag == "tbl":
            parent.insert(index, deepcopy(block))
            index += 1

        elif tag == "p":
            rId = _extract_blip_rId(block)
            if rId:
                img_path = _save_image_from_part(source_doc, rId)
                if img_path:
                    p = target_doc.add_paragraph()
                    p.add_run().add_picture(img_path, width=Cm(14))
                    parent.insert(index, p._element)
                    index += 1
                    cleanup(img_path)
            else:
                parent.insert(index, deepcopy(block))
                index += 1

    log.info("Section %s–%s inserted at '%s'", start_heading, stop_heading, marker_text)


# ==========================================================
# TABLE-ONLY SECTION EXTRACTION
# ==========================================================
def insert_section_tables_only(target_doc, source_doc, marker_text: str,
                                start_heading: str, stop_heading: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    parent.remove(marker_el)
    capture = False

    for tag, block in _iter_body(source_doc):
        if tag == "p":
            text = _block_text(block)
            if start_heading in text:
                capture = True
                continue
            if capture and stop_heading and stop_heading in text:
                break

        if capture and tag == "tbl":
            parent.insert(index, deepcopy(block))
            index += 1


# ==========================================================
# FIRST TABLE AFTER HEADING
# ==========================================================
def copy_first_table_after_heading(source_doc, target_doc, marker_text: str,
                                    heading_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    found_heading = False

    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text in _block_text(block):
            found_heading = True
            continue
        if found_heading and tag == "tbl":
            parent.insert(index, deepcopy(block))
            parent.remove(marker_el)
            log.info("First table after '%s' inserted.", heading_text)
            return

    log.warning("No table found after heading '%s'.", heading_text)


# ==========================================================
# IMAGE HELPERS
# ==========================================================
def _extract_blip_rId(block) -> str | None:
    ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    for node in block.iter():
        if node.tag.endswith("blip"):
            return node.get(f"{{{ns}}}embed")
    return None


def _save_image_from_part(source_doc, rId: str) -> str | None:
    try:
        part = source_doc.part.related_parts.get(rId) or \
               source_doc.part.rels.get(rId, {}).target_part
        if part is None:
            return None
        img_path = os.path.join(app.config["UPLOAD_FOLDER"], f"tmp_{uuid.uuid4().hex}.png")
        with open(img_path, "wb") as f:
            f.write(part.blob)
        return img_path
    except Exception as e:
        log.error("Image extraction failed rId=%s: %s", rId, e)
        return None


def copy_first_image_after_main_heading(source_doc, target_doc, marker_text: str,
                                         heading_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    found = False

    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text.lower() in _block_text(block).lower():
            found = True
            continue
        if found and tag == "p":
            rId = _extract_blip_rId(block)
            if rId:
                img_path = _save_image_from_part(source_doc, rId)
                if img_path:
                    p = target_doc.add_paragraph()
                    p.add_run().add_picture(img_path)
                    parent.insert(index, p._element)
                    parent.remove(marker_el)
                    cleanup(img_path)
                    log.info("First image after '%s' inserted.", heading_text)
                    return

    log.warning("No image found after heading '%s'.", heading_text)


def copy_graph_after_table(source_doc, target_doc, marker_text: str,
                            heading_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    found_heading = False
    table_passed = False

    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text in _block_text(block):
            found_heading = True
            continue
        if found_heading:
            if tag == "tbl":
                table_passed = True
                continue
            if table_passed and tag == "p":
                rId = _extract_blip_rId(block)
                if rId:
                    img_path = _save_image_from_part(source_doc, rId)
                    if img_path:
                        p = target_doc.add_paragraph()
                        p.add_run().add_picture(img_path)
                        parent.insert(index, p._element)
                        parent.remove(marker_el)
                        cleanup(img_path)
                        log.info("Graph after table in '%s' inserted.", heading_text)
                        return

    log.warning("No graph found after table in section '%s'.", heading_text)


# ==========================================================
# RSA SUMMARY TABLE (merged tables)
# ==========================================================
from copy import deepcopy
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _fix_run_spacing(tbl_el):
    """Ensure all <w:t> elements preserve whitespace."""
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    for t_el in tbl_el.iter(f"{{{W}}}t"):
        text = t_el.text or ""
        if not text.strip() or text != text.strip() or text == " ":
            t_el.set(XML_SPACE, "preserve")


def _merge_cell_paragraphs(tbl_el):
    """
    Merge multiple paragraphs in a cell into one paragraph with line breaks.
    This fixes cells where text is split across paragraphs causing mid-word
    wrapping when the table is scaled to a narrower width.
    Skip the first row (header) as it is already handled by _fix_header_row_text.
    """
    rows = tbl_el.findall(f'{{{W}}}tr')
    if not rows:
        return

    for row in rows[1:]:  # skip header row
        for tc in row.findall(f'{{{W}}}tc'):
            paragraphs = tc.findall(f'{{{W}}}p')
            if len(paragraphs) <= 1:
                continue  # nothing to merge

            # Collect all runs from all paragraphs
            all_runs = []
            first_pPr = None
            for i, para in enumerate(paragraphs):
                if i == 0:
                    first_pPr = para.find(f'{{{W}}}pPr')
                runs = para.findall(f'{{{W}}}r')
                if i > 0 and runs:
                    # Add a line break run before runs from subsequent paragraphs
                    all_runs.append('BR')
                all_runs.extend(runs)

            # Remove all existing paragraphs
            for para in paragraphs:
                tc.remove(para)

            # Build one clean paragraph
            new_p = etree.SubElement(tc, f'{{{W}}}p')

            # Keep original paragraph properties from first paragraph
            if first_pPr is not None:
                new_p.insert(0, deepcopy(first_pPr))

            # Add all runs with line breaks between paragraphs
            for item in all_runs:
                if item == 'BR':
                    br_run = etree.SubElement(new_p, f'{{{W}}}r')
                    etree.SubElement(br_run, f'{{{W}}}br')
                else:
                    new_p.append(deepcopy(item))

def _remove_table_paragraph_style(tbl_el):
    """
    Remove w:pStyle references so cells use direct formatting only.
    This fixes rendering issues when the style hasn't been resolved
    at the insertion point in the target document.
    """
    for pPr in tbl_el.iter(f"{{{W}}}pPr"):
        for pStyle in pPr.findall(f"{{{W}}}pStyle"):
            pPr.remove(pStyle)
            
def _normalize_table_width(tbl_el, content_width):
    """Resize a table element (lxml) to fit the given content width."""
    tbl_pr = tbl_el.find(f'{{{W}}}tblPr')
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_el, f'{{{W}}}tblPr')

    tbl_w_el = tbl_pr.find(f'{{{W}}}tblW')
    if tbl_w_el is None:
        tbl_w_el = etree.SubElement(tbl_pr, f'{{{W}}}tblW')
    tbl_w_el.set(f'{{{W}}}w', str(content_width))
    tbl_w_el.set(f'{{{W}}}type', 'dxa')

    tbl_ind = tbl_pr.find(f'{{{W}}}tblInd')
    if tbl_ind is None:
        tbl_ind = etree.SubElement(tbl_pr, f'{{{W}}}tblInd')
    tbl_ind.set(f'{{{W}}}w', '0')
    tbl_ind.set(f'{{{W}}}type', 'dxa')

    scale = 1.0
    tbl_grid = tbl_el.find(f'{{{W}}}tblGrid')
    if tbl_grid is not None:
        grid_cols = tbl_grid.findall(f'{{{W}}}gridCol')
        if grid_cols:
            original_total = sum(
                int(c.get(f'{{{W}}}w', 0)) for c in grid_cols
            )
            if original_total > 0 and original_total != content_width:
                scale = content_width / original_total
                for col in grid_cols:
                    orig = int(col.get(f'{{{W}}}w', 0))
                    col.set(f'{{{W}}}w', str(int(orig * scale)))

    for tc in tbl_el.iter(f'{{{W}}}tc'):
        tc_pr = tc.find(f'{{{W}}}tcPr')
        if tc_pr is None:
            continue
        tc_w = tc_pr.find(f'{{{W}}}tcW')
        if tc_w is not None:
            w_type = tc_w.get(f'{{{W}}}type', 'dxa')
            if w_type == 'dxa':
                orig = int(tc_w.get(f'{{{W}}}w', 0))
                if orig > 0:
                    tc_w.set(f'{{{W}}}w', str(int(orig * scale)))


def _get_content_width_at_marker(target_doc, marker_el):
    """
    Find the effective portrait page content width at the marker's position.
    Skips landscape sections (used for wide data tables) and returns
    the most recent portrait sectPr before the marker.
    """
    body = target_doc.element.body
    all_elements = list(body)

    best_page_w = 11906
    best_margin_left = 1440
    best_margin_right = 1440

    def _is_landscape(sect_pr):
        pg_sz = sect_pr.find(f'{{{W}}}pgSz')
        if pg_sz is None:
            return False
        orient = pg_sz.get(f'{{{W}}}orient', '')
        if orient == 'landscape':
            return True
        w = int(pg_sz.get(f'{{{W}}}w', 0))
        h = int(pg_sz.get(f'{{{W}}}h', 1))
        return w > h

    # Load document-level sectPr as baseline
    doc_sect_pr = body.find(f'{{{W}}}sectPr')
    if doc_sect_pr is not None and not _is_landscape(doc_sect_pr):
        pg_sz = doc_sect_pr.find(f'{{{W}}}pgSz')
        pg_mar = doc_sect_pr.find(f'{{{W}}}pgMar')
        if pg_sz is not None:
            best_page_w = int(pg_sz.get(f'{{{W}}}w', best_page_w))
        if pg_mar is not None:
            best_margin_left = int(pg_mar.get(f'{{{W}}}left', best_margin_left))
            best_margin_right = int(pg_mar.get(f'{{{W}}}right', best_margin_right))

    # Walk elements, only update from portrait sections
    for el in all_elements:
        if el is marker_el:
            break
        sect_pr = el.find(f'.//{{{W}}}sectPr')
        if sect_pr is not None and not _is_landscape(sect_pr):
            pg_sz = sect_pr.find(f'{{{W}}}pgSz')
            pg_mar = sect_pr.find(f'{{{W}}}pgMar')
            if pg_sz is not None:
                best_page_w = int(pg_sz.get(f'{{{W}}}w', best_page_w))
            if pg_mar is not None:
                best_margin_left = int(pg_mar.get(f'{{{W}}}left', best_margin_left))
                best_margin_right = int(pg_mar.get(f'{{{W}}}right', best_margin_right))

    return best_page_w - best_margin_left - best_margin_right


def _fix_header_row_text(tbl_el):
    """Replace the first row's cell content with clean plain text."""
    rows = tbl_el.findall(f'{{{W}}}tr')
    if not rows:
        return

    header_row = rows[0]

    for tc in header_row.findall(f'{{{W}}}tc'):
        for p in tc.findall(f'{{{W}}}p'):
            tc.remove(p)

        new_p = etree.SubElement(tc, f'{{{W}}}p')

        pPr = etree.SubElement(new_p, f'{{{W}}}pPr')
        jc = etree.SubElement(pPr, f'{{{W}}}jc')
        jc.set(f'{{{W}}}val', 'center')
        rPr_p = etree.SubElement(pPr, f'{{{W}}}rPr')
        etree.SubElement(rPr_p, f'{{{W}}}b')
        sz = etree.SubElement(rPr_p, f'{{{W}}}sz')
        sz.set(f'{{{W}}}val', '24')
        fonts = etree.SubElement(rPr_p, f'{{{W}}}rFonts')
        fonts.set(f'{{{W}}}ascii', 'Arial')

        lines = [
            "ROAD SAFETY AUDITOR RECOMMENDATION FOR ROAD SIGNAGES",
            "DURING",
            "OPERATION AND MAINTENANCE STAGE",
        ]

        for i, line in enumerate(lines):
            run = etree.SubElement(new_p, f'{{{W}}}r')
            rPr = etree.SubElement(run, f'{{{W}}}rPr')
            etree.SubElement(rPr, f'{{{W}}}b')
            sz2 = etree.SubElement(rPr, f'{{{W}}}sz')
            sz2.set(f'{{{W}}}val', '24')
            fonts2 = etree.SubElement(rPr, f'{{{W}}}rFonts')
            fonts2.set(f'{{{W}}}ascii', 'Arial')
            t = etree.SubElement(run, f'{{{W}}}t')
            t.text = line
            if i < len(lines) - 1:
                br_run = etree.SubElement(new_p, f'{{{W}}}r')
                etree.SubElement(br_run, f'{{{W}}}br')

def _fix_cell_indentation(tbl_el):
    """Remove indent and line spacing from all body cell paragraphs."""
    rows = tbl_el.findall(f'{{{W}}}tr')
    if not rows:
        return
    for row in rows[1:]:  # skip header row
        for tc in row.findall(f'{{{W}}}tc'):
            for para in tc.findall(f'{{{W}}}p'):
                pPr = para.find(f'{{{W}}}pPr')
                if pPr is None:
                    continue
                # Remove indentation
                for ind in pPr.findall(f'{{{W}}}ind'):
                    pPr.remove(ind)
                # Remove line spacing overrides
                for spacing in pPr.findall(f'{{{W}}}spacing'):
                    pPr.remove(spacing)
                # Also strip w:spacing from every run's rPr in this para
                for run in para.findall(f'{{{W}}}r'):
                    rPr = run.find(f'{{{W}}}rPr')
                    if rPr is not None:
                        for sp in rPr.findall(f'{{{W}}}spacing'):
                            rPr.remove(sp)


def insert_rsa_summary_table(source_doc, target_doc, marker_text: str):
    tables = source_doc.tables
    if not tables:
        log.warning("No tables in source doc for RSA summary.")
        return

    # Collect consecutive tables from the end with same column count
    merged = [tables[-1]]
    for i in range(len(tables) - 2, -1, -1):
        if len(tables[i].columns) == len(merged[0].columns):
            merged.insert(0, tables[i])
        else:
            break

    log.info("RSA summary: merging %d table(s).", len(merged))

    base_el = deepcopy(merged[0]._element)
    for tbl in merged[1:]:
        for row in tbl.rows:
            base_el.append(deepcopy(row._element))
    
    # Get marker first so we can measure the correct section width
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    # Compute width at the actual insertion point
    content_width = _get_content_width_at_marker(target_doc, marker_el)
    log.info(
        "RSA summary: content width at insertion = %d DXA (%.2f inches)",
        content_width,
        content_width / 1440,
    )

    
    

    _remove_table_paragraph_style(base_el)
    _fix_run_spacing(base_el)
    _fix_cell_indentation(base_el)
    _merge_cell_paragraphs(base_el)
    _normalize_table_width(base_el, content_width)
    _fix_header_row_text(base_el)
    
    parent.insert(index, base_el)
    parent.remove(marker_el)
    log.info("RSA summary table inserted.")
# ==========================================================
# FULL DOCUMENT INSERT  (Annexures A/B/C)

# ==========================================================
def insert_full_document(source_doc, target_doc, marker_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    for element in source_doc.element.body:
        parent.insert(index, deepcopy(element))
        index += 1

    parent.remove(marker_el)
    log.info("Full document inserted at '%s'.", marker_text)


# ==========================================================
# ANNEXURE D — extract from heading to end
# ==========================================================
def extract_till_end(source_doc, start_heading: str) -> list:
    content = []
    capture = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and start_heading in _block_text(block):
            capture = True
            continue
        if capture:
            content.append(deepcopy(block))
    return content


def insert_section_blocks(doc, marker_text: str, content_blocks: list,source_doc):
    parent, index, marker_el = _find_marker(doc, marker_text)
    if parent is None:
        return

    for block in content_blocks:
        tag = block.tag.split("}")[-1]

        if tag == "p":
            new_para = doc.add_paragraph()
            for node in block.iter():
                if node.tag.endswith("t") and node.text:
                    new_para.add_run(node.text)
                if node.tag.endswith("blip"):
                    rId = node.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if rId and rId in source_doc.part.rels:
                        img_path = _save_image_from_part(source_doc, rId)
                        if img_path:
                            run = new_para.add_run()
                            run.add_picture(img_path, width=Cm(14.4))
                            cleanup(img_path)

            parent.insert(index, new_para._element)
            index += 1

        elif tag == "tbl":
            tbl_copy = deepcopy(block)
            parent.insert(index, tbl_copy)
            index += 1
            # Apply formatting to the just-inserted table
            table_obj = doc.tables[-1]
            format_table(table_obj)

    parent.remove(marker_el)
    log.info("Section blocks inserted at '%s' (%d blocks).", marker_text, len(content_blocks))


# ==========================================================
# ZONE COLOR
# ==========================================================
def apply_zone_color(doc, zone: str):
    color = ZONE_COLORS.get(zone)
    if not color:
        log.warning("Unknown zone: %s", zone)
        return
    for para in doc.paragraphs:
        for run in para.runs:
            if zone in run.text:
                run.font.color.rgb = color
                run.bold = True
    log.info("Zone color applied for %s.", zone)


# ==========================================================
# ROAD DATA PARSING
# ==========================================================
def parse_road_data(form) -> tuple[dict, list]:
    """
    Returns:
        road_data    = {"Service Roads": {"LHS": 2, "RHS": 3}, ...}
        road_summary = [{"name": "Service Roads", "total": 5}, ...]
    """
    selected_types = form.getlist("road_type")
    road_data = {}
    road_summary = []

    for road_type in selected_types:
        lhs = int(form.get(f"{road_type}_LHS", 0) or 0)
        rhs = int(form.get(f"{road_type}_RHS", 0) or 0)
        road_data[road_type] = {"LHS": lhs, "RHS": rhs}
        total = lhs + rhs
        if total > 0:
            road_summary.append({"name": road_type, "total": total})

    return road_data, road_summary


# ==========================================================
# MAIN ROUTE
# ==========================================================
@app.route("/", methods=["GET", "POST"])
def form():
    if request.method != "POST":
        return render_template("form.html")

    ensure_folders()

    temp_files = []   # track all temp paths for cleanup

    try:
        template_path = app.config["TEMPLATE_PATH"]
        tpl = DocxTemplate(template_path)

        # ── DATE ──────────────────────────────────────────
        raw_date = request.form.get("starting_survey_date", "")
        try:
            formatted_date = datetime.strptime(raw_date, "%Y-%m-%d").strftime("%d %B %Y")
        except ValueError:
            formatted_date = raw_date

        # ── MAP IMAGE ─────────────────────────────────────
        image = None
        image_file = request.files.get("map_image")
        if image_file and image_file.filename:
            image_path = save_upload(image_file, "map")
            temp_files.append(image_path)
            image = InlineImage(tpl, image_path, width=Cm(12))

        # ── ROAD DATA ─────────────────────────────────────
        road_data, road_summary = parse_road_data(request.form)

        # ── RENDER TEMPLATE ───────────────────────────────
        temp_path = os.path.join(app.config["UPLOAD_FOLDER"], f"rendered_{uuid.uuid4().hex}.docx")
        temp_files.append(temp_path)

        context = {
            "project_name":       request.form.get("project_name", ""),
            "upc_code":           request.form.get("upc_code", ""),
            "state":              request.form.get("state", ""),
            "ro":                 request.form.get("ro", ""),
            "piu":                request.form.get("piu", ""),
            "length":             request.form.get("length", ""),
            "flexibleorrigid":    request.form.get("flexibleorrigid", ""),
            "lanes":              request.form.get("lanes", ""),
            "om_dlp":             request.form.get("om_dlp", ""),
            "starting_survey_date": formatted_date,
            "zone":               request.form.get("zone", ""),
            "map_image":          image,
            "road_summary":       road_summary,
        }

        tpl.render(context)
        tpl.save(temp_path)

        # ── OPEN RENDERED DOC FOR MUTATIONS ───────────────
        doc = Document(temp_path)

        # ── ANALYSED DOCUMENT ─────────────────────────────
        analysed_file = request.files.get("analysed_doc")
        if analysed_file and analysed_file.filename:
            analysed_path = save_upload(analysed_file, "analysed")
            temp_files.append(analysed_path)
            source_doc = Document(analysed_path)

            # Annexure D
            anx_d_blocks = extract_till_end(source_doc, "Chainage Wise Gap Analysis")
            insert_section_blocks(doc, "### ANX_D ###", anx_d_blocks, source_doc)

            # Executive summary table
            copy_executive_summary_table(source_doc, doc,
                                          "### INSERT_EXECUTIVE_SUMMARY_TABLE ###")

            # Section 3.1
            insert_section(doc, source_doc,
                           "### INSERT_INVENTORY_SECTION_ONE ###", "3.1", "3.2")

            # Section 3.2
            insert_section(doc, source_doc,
                           "### INSERT_INVENTORY_SECTION_TWO ###", "3.2", "3.3")

            # Section 4.1 (full)
            insert_section(doc, source_doc,
                           "### RESULT_GAP_STUDY_ONE ###", "4.1", "4.2")

            # Section 4.2 (tables only)
            insert_section_tables_only(doc, source_doc,
                                        "### RESULT_GAP_STUDY_TWO ###", "4.2", "4.3")

            # Section 4.2 — first table
            copy_first_table_after_heading(source_doc, doc,
                                            "### RESULT_GAP_STUDY_TWO_TABLE ###", "4.2")

            # Section 4.2 — first graph
            copy_first_image_after_main_heading(source_doc, doc,
                                                 "### RESULT_GAP_STUDY_TWO_GRAPH_ONE ###",
                                                 "Chainage Wise Gap Analysis")

            # Section 4.2 — second graph (after table)
            copy_graph_after_table(source_doc, doc,
                                    "### RESULT_GAP_STUDY_TWO_GRAPH_TWO ###", "4.2")

        # ── ANNEXURE A ────────────────────────────────────
        anx_a = request.files.get("anx_a_doc")
        if anx_a and anx_a.filename:
            p = save_upload(anx_a, "anx_a")
            temp_files.append(p)
            insert_full_document(Document(p), doc, "### ANX_A ###")

        # ── ANNEXURE B ────────────────────────────────────
        anx_b = request.files.get("anx_b_doc")
        if anx_b and anx_b.filename:
            p = save_upload(anx_b, "anx_b")
            temp_files.append(p)
            insert_full_document(Document(p), doc, "### ANX_B ###")

        # ── ANNEXURE C ────────────────────────────────────
        anx_c = request.files.get("anx_c_doc")
        if anx_c and anx_c.filename:
            p = save_upload(anx_c, "anx_c")
            temp_files.append(p)
            anx_c_doc = Document(p)
            insert_full_document(anx_c_doc, doc, "### ANX_C ###")
            insert_rsa_summary_table(anx_c_doc, doc, "### RSA_SUMMARY ###")

        # ── ROAD TABLE ────────────────────────────────────
        if road_data:
            insert_road_table(doc, "### INSERT_ROAD_TABLE ###", road_data)

        # ── ZONE COLOR ────────────────────────────────────
        apply_zone_color(doc, request.form.get("zone", ""))

        # ── SAVE OUTPUT ───────────────────────────────────
        project_name = request.form.get("project_name", "Report")
        safe_name = sanitize_filename(f"Final_Report_{project_name}")
        final_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{safe_name}.docx")
        doc.save(final_path)
        log.info("Report saved: %s", final_path)

        # ── CLEANUP AFTER RESPONSE ────────────────────────
        @after_this_request
        def cleanup_handler(response):
            cleanup(*temp_files)
            return response

        return send_file(final_path, as_attachment=True,
                         download_name=f"{safe_name}.docx")

    except Exception as e:
        log.exception("Report generation failed: %s", e)
        cleanup(*temp_files)
        return f"<h3>Error generating report</h3><pre>{e}</pre>", 500


if __name__ == "__main__":
    app.run(debug=True)
